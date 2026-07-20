"""Reports API — create & retrieve competitive analysis reports."""

from __future__ import annotations

import asyncio
from datetime import datetime
from uuid import UUID

from fastapi import APIRouter, HTTPException

from app.application.dto.report_dto import (
    ReportCreateRequest,
    ReportCreateResponse,
    ReportDetailResponse,
    ReportListResponse,
    ReportSectionDTO,
)
from app.infrastructure.workflow.state import create_initial_state
from app.infrastructure.trace import trace_collector, TraceStatus

router = APIRouter(prefix="/reports", tags=["reports"])

# ── File-backed persistence ──
from app.infrastructure.persistence.file_store import (
    load_reports,
    save_reports,
    load_tasks,
    save_tasks,
)

_reports: dict[str, dict] = load_reports()
_tasks: dict[str, dict] = load_tasks()


# ── Persistence helpers ──

def _persist_reports() -> None:
    """Write _reports to disk synchronously (called from sync context)."""
    save_reports(dict(_reports))


def _persist_tasks() -> None:
    """Write _tasks to disk synchronously."""
    save_tasks(dict(_tasks))


def _diagnose_and_attach(task_id_str: str, exc: Exception) -> dict:
    """Run diagnosis on a failed task and attach result to _tasks state."""
    try:
        from app.infrastructure.trace.diagnosis import diagnose_task
        diag = diagnose_task(task_id_str, trace_collector)
        return diag.to_dict()
    except Exception:
        return {
            "task_id": task_id_str,
            "error_type": "LLM_ERROR",
            "root_cause": f"{type(exc).__name__}: {str(exc)[:200]}",
            "suggestion": "请检查后台日志获取详细错误信息",
            "retry_available": True,
        }


@router.post("", response_model=ReportCreateResponse)
async def create_report(body: ReportCreateRequest) -> ReportCreateResponse:
    """Create a new competitive analysis report.

    Phase 2: Returns immediately. Workflow runs in background.
    """
    from app.infrastructure.workflow.graph import workflow_graph
    from app.infrastructure.workflow.stream import ensure_listener

    state = create_initial_state(body.model_dump())
    task_id_str = str(state["task_id"])

    # --- Trace: API request ---
    trace_collector.record_trace(
        task_id=task_id_str,
        stage="api",
        agent_name="",
        status=TraceStatus.SUCCESS,
        input_summary=f"POST /api/reports our={body.our_company}, competitor={body.competitor_company}, product={body.product}",
        output_summary=f"task_id={task_id_str}",
        metadata={"method": "POST", "path": "/api/reports", "params": body.model_dump()},
    )
    # ---

    _tasks[task_id_str] = {
        "task_id": task_id_str,
        "status": "pending",
        "state": state,
    }

    _persist_tasks()

    ensure_listener(task_id_str)

    async def _run_workflow():
        try:
            # --- Trace: workflow start ---
            wf_trace = trace_collector.start_trace(
                task_id=task_id_str,
                stage="workflow",
                agent_name="",
                input_summary=f"company={body.our_company}, competitor={body.competitor_company}",
            )
            # ---
            final_state = state
            # Phase tracking — accumulate phase_history from node outputs
            _phase_history = []
            _phase_node_map = {
                "validate_input_node": "validated",
                "plan_node": "planned",
                "research_node": "researched",
                "compare_node": "compared",
                "insight_node": "insighted",
                "strategy_node": "strategized",
                "report_node": "reported",
                "review_node": "reviewed",
            }
            async for chunk in workflow_graph.astream(state, stream_mode="updates"):
                for node_name, node_state in chunk.items():
                    final_state.update(node_state)
                    # Track phase progress
                    phase_name = _phase_node_map.get(node_name)
                    if phase_name:
                        now = datetime.utcnow().isoformat()
                        _phase_history.append({
                            "phase": phase_name,
                            "entered_at": now,
                            "duration_ms": 0,
                            "status": "completed",
                        })
                        final_state["phase_history"] = list(_phase_history)
                    # Update progress based on phase
                    phase_progress_map = {
                        "validated": 5, "planned": 15, "researched": 40,
                        "compared": 55, "insighted": 60, "strategized": 65,
                        "reported": 85, "reviewed": 95,
                    }
                    if phase_name:
                        final_state["progress"] = float(phase_progress_map.get(phase_name, 0))
                        final_state["current_agent"] = phase_name
                    _tasks[task_id_str]["state"] = dict(final_state)
            final_phase = final_state.get("current_phase", "")
            if "fail" in final_phase or final_phase == "validation_failed":
                _tasks[task_id_str]["status"] = "failed"
                _tasks[task_id_str]["error"] = (
                    "; ".join(str(e) for e in final_state.get("errors", []) if e)
                    or "分析流程失败"
                )
                # Attach diagnosis on workflow failure
                _tasks[task_id_str]["diagnosis"] = _diagnose_and_attach(
                    task_id_str, Exception(_tasks[task_id_str]["error"])
                )
            else:
                _tasks[task_id_str]["status"] = "completed"

            _persist_tasks()

            report_doc = final_state.get("report_document") or {}
            sections_data = report_doc.get("sections", []) if isinstance(report_doc, dict) else []

            # --- Trace: result (after workflow completion) ---
            trace_collector.record_trace(
                task_id=task_id_str,
                stage="result",
                agent_name="",
                status=TraceStatus.SUCCESS if _tasks[task_id_str]["status"] != "failed" else TraceStatus.FAILED,
                input_summary=f"final_phase={final_state.get('current_phase', 'unknown')}",
                output_summary=f"sections={len(sections_data)}, words={(report_doc or {}).get('metadata', {}).get('total_word_count', 0)}",
                metadata={
                    "section_count": len(sections_data),
                    "word_count": (report_doc or {}).get("metadata", {}).get("total_word_count", 0),
                    "final_phase": final_state.get("current_phase", ""),
                    "errors": final_state.get("errors", []),
                },
            )
            # ---

            _reports[task_id_str] = ReportDetailResponse(
                id=UUID(final_state["task_id"]),
                task_id=UUID(final_state["task_id"]),
                our_company=body.our_company,
                competitor_company=body.competitor_company,
                product=body.product,
                objective=body.scene or body.objective,
                markdown=(report_doc or {}).get("formats", {}).get("markdown"),
                html=(report_doc or {}).get("formats", {}).get("html"),
                # Convert filesystem path to API URL for frontend
                word_url=f"/api/reports/{task_id_str}/download" if (report_doc or {}).get("formats", {}).get("docx_url") else None,
                sections=[ReportSectionDTO(**s) for s in sections_data],
                total_word_count=(report_doc or {}).get("metadata", {}).get("total_word_count", 0),
                created_at=datetime.utcnow(),
            ).model_dump()
        except Exception as exc:
            import traceback
            error_msg = f"{type(exc).__name__}: {exc}"
            _tasks[task_id_str]["status"] = "failed"
            _tasks[task_id_str]["error"] = error_msg
            if _tasks[task_id_str].get("state"):
                _tasks[task_id_str]["state"]["error_info"] = error_msg

            # --- Auto-diagnose on exception ---
            diagnosis = _diagnose_and_attach(task_id_str, exc)
            _tasks[task_id_str]["diagnosis"] = diagnosis
            # ---

            _persist_tasks()

            # Trace the workflow-level failure
            trace_collector.record_trace(
                task_id=task_id_str,
                stage="workflow",
                agent_name="",
                status=TraceStatus.FAILED,
                error=error_msg,
                output_summary=f"failed: {error_msg[:300]}",
                metadata={"diagnosis": diagnosis},
            )

            print(f"[WORKFLOW ERROR] task={task_id_str}: {error_msg}")
            traceback.print_exc()

            _reports[task_id_str] = {
                "id": task_id_str,
                "task_id": task_id_str,
                "our_company": body.our_company,
                "competitor_company": body.competitor_company,
                "product": body.product,
                "objective": body.scene or body.objective,
                "status": "failed",
                "error": error_msg,
                "diagnosis": diagnosis,
                "markdown": None,
                "html": None,
                "word_url": None,
                "sections": [],
                "total_word_count": 0,
                "created_at": datetime.utcnow().isoformat(),
            }
        finally:
            _persist_reports()

    asyncio.create_task(_run_workflow())

    return ReportCreateResponse(
        task_id=state["task_id"],
        status="pending",
        message="分析任务已创建",
    )


@router.get("/{task_id}", response_model=ReportDetailResponse)
async def get_report(task_id: UUID) -> ReportDetailResponse:
    """Get the final report by task ID.

    Returns 404 only when the task genuinely doesn't exist.
    For tasks still in progress, returns 202 with a status message.
    """
    report = _reports.get(str(task_id))
    if not report:
        # Check if task exists (but report not ready yet or workflow running)
        if str(task_id) in _tasks:
            raise HTTPException(
                status_code=202,
                detail="报告正在生成中，请稍后重试",
            )
        raise HTTPException(status_code=404, detail="报告不存在")
    return ReportDetailResponse(**report)



@router.get("/{task_id}/download")
async def download_report(task_id: UUID):
    """Download the Word report — generate DOCX from markdown on-the-fly."""
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    from docx import Document

    report = _reports.get(str(task_id))
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    markdown = report.get("markdown", "")
    if not markdown:
        raise HTTPException(status_code=404, detail="Word文件不存在或尚未生成")

    # Generate DOCX from markdown using python-docx
    try:
        doc = Document()
        # Use default font
        for line in markdown.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|"):
                continue  # Skip tables for now
            elif line == "---":
                doc.add_paragraph("_" * 40)
            else:
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Word文件生成失败: {e}")

    filename = f"竞品分析报告_{task_id[:8]}.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.delete("/{task_id}")
async def delete_report(task_id: UUID) -> dict:
    """Delete a report and its task data."""
    key = str(task_id)
    found = False
    if key in _reports:
        del _reports[key]
        found = True
    if key in _tasks:
        del _tasks[key]
        found = True
    if not found:
        raise HTTPException(status_code=404, detail="报告不存在")
    _persist_reports()
    _persist_tasks()
    return {"status": "deleted", "task_id": key}


@router.get("", response_model=ReportListResponse)
async def list_reports() -> ReportListResponse:
    """List all generated reports."""
    reports = list(_reports.values())
    return ReportListResponse(reports=reports, total=len(reports))
