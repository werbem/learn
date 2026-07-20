"""Report export utilities — Markdown, HTML, Word generation.

Not analysis — only formatting/layout.
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Optional

# python-docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor


# ═════════════════════════════════════════════════
#  Markdown Builder
# ═════════════════════════════════════════════════

class MarkdownBuilder:
    """Builds a Markdown string section by section."""

    def __init__(self) -> None:
        self._parts: list[str] = []

    def h1(self, text: str) -> None:
        self._parts.append(f"# {text}\n")

    def h2(self, text: str) -> None:
        self._parts.append(f"\n## {text}\n")

    def h3(self, text: str) -> None:
        self._parts.append(f"\n### {text}\n")

    def h4(self, text: str) -> None:
        self._parts.append(f"\n#### {text}\n")

    def para(self, text: str) -> None:
        self._parts.append(f"{text}\n")

    def quote(self, text: str) -> None:
        self._parts.append(f"> {text}\n\n")

    def code(self, text: str, lang: str = "") -> None:
        self._parts.append(f"```{lang}\n{text}\n```\n")

    def bullet(self, items: list[str]) -> None:
        for item in items:
            self._parts.append(f"- {item}\n")
        self._parts.append("\n")

    def bullets(self, items: list[str]) -> None:
        self.bullet(items)

    def table(self, headers: list[str], rows: list[list[str]]) -> None:
        """Generate a Markdown table."""
        sep = "| " + " | ".join("---" for _ in headers) + " |\n"
        self._parts.append("| " + " | ".join(headers) + " |\n")
        self._parts.append(sep)
        for row in rows:
            self._parts.append("| " + " | ".join(row) + " |\n")
        self._parts.append("\n")

    def image_placeholder(self, caption: str, width: str = "600px") -> None:
        """Image placeholder — add actual image in Phase 3."""
        self._parts.append(
            f'\n![{caption}](placeholder-{caption}.png "{caption}")\n\n'
        )

    def separator(self) -> None:
        self._parts.append("\n---\n")

    def build(self) -> str:
        return "".join(self._parts)


# ═════════════════════════════════════════════════
#  HTML Builder
# ═════════════════════════════════════════════════

class HTMLBuilder:
    """Builds a self-contained HTML page with inline CSS."""

    _CSS = """
    body { font-family: -apple-system, 'PingFang SC', 'Microsoft YaHei', sans-serif;
           max-width: 960px; margin: 0 auto; padding: 40px 20px;
           line-height: 1.8; color: #333; font-size: 15px; }
    h1 { font-size: 28px; border-bottom: 3px solid #2563eb; padding-bottom: 12px; margin-top: 40px; }
    h2 { font-size: 22px; border-bottom: 1px solid #e5e7eb; padding-bottom: 8px; margin-top: 36px; color: #1f2937; }
    h3 { font-size: 18px; margin-top: 28px; color: #374151; }
    h4 { font-size: 16px; margin-top: 20px; color: #4b5563; }
    table { border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 14px; }
    th { background: #2563eb; color: white; padding: 10px 12px; text-align: left; }
    td { border: 1px solid #d1d5db; padding: 8px 12px; vertical-align: top; }
    tr:nth-child(even) { background: #f9fafb; }
    blockquote { border-left: 4px solid #2563eb; margin: 16px 0; padding: 12px 20px;
                 background: #eff6ff; color: #1e40af; }
    code { background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-size: 13px; }
    pre { background: #1f2937; color: #f9fafb; padding: 16px; border-radius: 8px; overflow-x: auto; }
    pre code { background: none; color: inherit; padding: 0; }
    .cover { text-align: center; padding: 80px 0 40px; }
    .cover h1 { border: none; font-size: 32px; }
    .cover .meta { color: #6b7280; font-size: 14px; }
    .toc { background: #f9fafb; padding: 20px 24px; border-radius: 8px; margin: 20px 0; }
    .toc ul { list-style: none; padding-left: 0; }
    .toc li { padding: 4px 0; }
    .toc a { color: #2563eb; text-decoration: none; }
    .toc a:hover { text-decoration: underline; }
    .swot-cell { padding: 12px; min-height: 80px; }
    .swot-positive { background: #ecfdf5; }
    .swot-negative { background: #fef2f2; }
    .swot-header { background: #f3f4f6; font-weight: bold; text-align: center; }
    .footer { margin-top: 60px; padding-top: 20px; border-top: 1px solid #e5e7eb;
              color: #9ca3af; font-size: 12px; }
    img { max-width: 100%; border: 1px solid #e5e7eb; border-radius: 8px; margin: 12px 0; }
    """

    def __init__(self, title: str = "竞品分析报告") -> None:
        self._parts: list[str] = [
            "<!DOCTYPE html><html lang='zh-CN'><head><meta charset='UTF-8'>",
            f"<title>{title}</title>",
            f"<style>{self._CSS}</style></head><body>",
        ]

    def cover(self, our: str, competitor: str, product: str, date: str) -> None:
        self._parts.append(
            f"<div class='cover'>"
            f"<h1>互联网产品竞品分析报告</h1>"
            f"<p class='meta'>我方：{our} | 竞品：{competitor} | 产品：{product}</p>"
            f"<p class='meta'>生成日期：{date}</p>"
            f"</div>"
        )

    def toc(self, items: list[tuple[str, str]]) -> None:
        """items: [(section_id, section_title), ...]"""
        self._parts.append("<div class='toc'><h2>目录</h2><ul>")
        for sid, title in items:
            self._parts.append(f"<li><a href='#{sid}'>{title}</a></li>")
        self._parts.append("</ul></div>")

    def add_title(self, text: str) -> None:
        """Cover title block."""
        self._parts.append(f'<div class="cover"><h1>{text}</h1></div>')

    def h1(self, text: str) -> None:
        """H1 heading."""
        self._parts.append(f"<h1>{text}</h1>")

    def h2(self, text: str, sid: str = "") -> None:
        aid = f" id='{sid}'" if sid else ""
        self._parts.append(f"<h2{aid}>{text}</h2>")

    def h3(self, text: str) -> None:
        self._parts.append(f"<h3>{text}</h3>")

    def h4(self, text: str) -> None:
        self._parts.append(f"<h4>{text}</h4>")

    def para(self, text: str) -> None:
        self._parts.append(f"<p>{text}</p>")

    def quote(self, text: str) -> None:
        self._parts.append(f"<blockquote>{text}</blockquote>")

    def code(self, text: str) -> None:
        self._parts.append(f"<pre><code>{text}</code></pre>")

    def bullet(self, items: list[str]) -> None:
        self._parts.append("<ul>")
        for item in items:
            self._parts.append(f"<li>{item}</li>")
        self._parts.append("</ul>")

    def bullets(self, items: list[str]) -> None:
        self.bullet(items)

    def table(self, headers: list[str], rows: list[list[str]]) -> None:
        self._parts.append("<table><thead><tr>")
        for h in headers:
            self._parts.append(f"<th>{h}</th>")
        self._parts.append("</tr></thead><tbody>")
        for row in rows:
            self._parts.append("<tr>")
            for cell in row:
                self._parts.append(f"<td>{cell}</td>")
            self._parts.append("</tr>")
        self._parts.append("</tbody></table>")

    def image_placeholder(self, caption: str) -> None:
        self._parts.append(
            f"<div style='text-align:center;padding:20px;background:#f9fafb;"
            f"border:2px dashed #d1d5db;border-radius:8px;margin:16px 0;'>"
            f"<p style='color:#9ca3af;'>📷 {caption}</p></div>"
        )

    def separator(self) -> None:
        self._parts.append("<hr>")

    def build(self) -> str:
        self._parts.append("<div class='footer'>")
        self._parts.append(
            f"<p>报告由 AI 竞品分析助手自动生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>"
        )
        self._parts.append("</div></body></html>")
        return "".join(self._parts)


# ═════════════════════════════════════════════════
#  Word Builder  (python-docx)
# ═════════════════════════════════════════════════

class WordBuilder:
    """Builds a .docx document with proper formatting."""

    def __init__(self) -> None:
        self.doc = Document()

        # Page setup
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        # Style setup
        self._setup_styles()

    def _setup_styles(self) -> None:
        """Configure built-in styles for Chinese document."""
        style = self.doc.styles['Normal']
        style.font.name = 'PingFang SC'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(4)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

        for level, size, color in [
            ('Heading 1', 22, '1F2937'),
            ('Heading 2', 18, '2563EB'),
            ('Heading 3', 15, '374151'),
        ]:
            hs = self.doc.styles[level]
            hs.font.name = 'PingFang SC'
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16),
            )
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

    def add_title(self, text: str) -> None:
        """Document cover title."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        p.paragraph_format.space_before = Pt(80)
        p.paragraph_format.space_after = Pt(20)

    def add_meta(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        p.paragraph_format.space_after = Pt(4)

    def add_h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)

    def add_h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def add_h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def add_para(self, text: str) -> None:
        self.doc.add_paragraph(text)

    def add_quote(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        pf = p.paragraph_format
        pf.left_indent = Inches(0.4)
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        # Add left border via XML
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr %s>'
            '<w:left w:val="single" w:sz="12" w:space="8" w:color="2563EB"/>'
            '</w:pBdr>' % nsdecls('w')
        )
        pPr.append(pBdr)

    def add_bullets(self, items: list[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
    ) -> None:
        """Add a formatted table."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            # Blue header background
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="2563EB" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True

        # Data rows
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                table.rows[ri + 1].cells[ci].text = val
                # Alternating row color
                if ri % 2 == 1:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F9FAFB" w:val="clear"/>'
                    )
                    table.rows[ri + 1].cells[ci]._tc.get_or_add_tcPr().append(shading)

        # Add spacing after table
        self.doc.add_paragraph()

    def add_image_placeholder(self, caption: str) -> None:
        """Placeholder for image — real images in Phase 3."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        # Create a bordered box
        run = p.add_run(f"[图表占位] {caption}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    def page_break(self) -> None:
        self.doc.add_page_break()


    # ── Method aliases for builder consistency ──
    def para(self, text: str) -> None:
        self.add_para(text)

    def h1(self, text: str) -> None:
        self.add_h1(text)

    def quote(self, text: str) -> None:
        self.add_quote(text)

    def bullets(self, items: list[str]) -> None:
        self.add_bullets(items)

    def table(self, headers: list[str], rows: list[list[str]]) -> None:
        self.add_table(headers, rows)

    def image_placeholder(self, caption: str) -> None:
        self.add_image_placeholder(caption)

    def save(self, path: str) -> str:
        """Save document and return the file path."""
        os.makedirs(os.path.dirname(path), exist_ok=True)
        self.doc.save(path)
        return path
